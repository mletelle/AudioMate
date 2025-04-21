from docx import Document
from io import BytesIO

def test_generar_docx():
    doc = Document()
    doc.add_heading("Transcripción", 0)
    doc.add_paragraph("Archivo: test.mp3")
    doc.add_paragraph("Fecha: 2025-04-21")
    doc.add_paragraph("")

    segmentos = [
        {"start": 0.0, "end": 2.3, "text": "Hola, ¿cómo estás?"},
        {"start": 2.4, "end": 4.0, "text": "Todo bien por acá."}
    ]

    for seg in segmentos:
        doc.add_paragraph(f"[{seg['start']:.2f}s - {seg['end']:.2f}s] {seg['text']}", style='List Bullet')

    buffer = BytesIO()
    doc.save(buffer)

    assert buffer.tell() > 0, "El archivo docx no se generó correctamente"
