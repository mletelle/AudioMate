#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AudioMate – Transcripción de audio a texto
-----------------------------------------
- Procesa archivos de audio (MP3, WAV, M4A) de hasta 1 GB.
- Muestra el progreso de la transcripción en tiempo real.
- Utiliza GPU para acelerar el procesamiento si está disponible.
- Permite la descarga de transcripciones en formatos TXT, DOCX y ZIP.
"""

import io
import json
import re
import logging
import os
import subprocess
import zipfile
import time
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any

import streamlit as st
import torch
import psutil
from pynvml import *
from docx import Document
from faster_whisper import WhisperModel

# --- Configuración General ---

UPLOAD_DIR = "uploads"
MAX_FILE_SIZE = 1_000 * 1024 * 1024  # 1000 MB
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Configuración del sistema de registro de eventos (logging)
logging.basicConfig(
    filename="audiomate.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)

st.set_page_config(page_title="AudioMate", layout="centered")
st.title("AudioMate – Transcripción de Audio a Texto")

logging.info("Aplicación AudioMate iniciada.")

# --- Detección de Dispositivo y Configuración de Modelo ---

# Determina si se utilizará CPU o GPU (CUDA)
DEVICE = (
    "cpu"
    if os.getenv("WHISPER_FORCE_CPU") == "1"
    else ("cuda" if torch.cuda.is_available() else "cpu")
)
st.sidebar.info(f"Dispositivo en uso: **{'GPU' if DEVICE == 'cuda' else 'CPU'}**")
logging.info(f"Dispositivo seleccionado: {DEVICE}")

# --- Monitor de Recursos del Sistema ---

def get_system_usage():
    """Obtiene el porcentaje de uso de CPU, RAM y GPU."""
    # Llamar a cpu_percent una vez con intervalo para inicializarlo correctamente
    psutil.cpu_percent(interval=None)
    usage = {
        "cpu": f"{psutil.cpu_percent(interval=None)}%",
        "ram": f"{psutil.virtual_memory().percent}%"
    }
    if DEVICE == "cuda":
        try:
            nvmlInit()
            handle = nvmlDeviceGetHandleByIndex(0)
            utilization = nvmlDeviceGetUtilizationRates(handle)
            usage["gpu"] = f"{utilization.gpu}%"
        except NVMLError as error:
            logging.error(f"Error al obtener métricas de NVML: {error}")
            usage["gpu"] = "Error"
        finally:
            nvmlShutdown()
    return usage

# Contenedor para mostrar las estadísticas que se actualizará periódicamente
usage_placeholder = st.sidebar.empty()
usage_stats = get_system_usage()
gpu_info = f" | GPU: **{usage_stats.get('gpu', 'N/A')}**" if DEVICE == 'cuda' else ""
usage_placeholder.info(f"Uso CPU: **{usage_stats['cpu']}** | RAM: **{usage_stats['ram']}**{gpu_info}")


# --- Selección y Carga del Modelo Whisper ---

# Selector para el modelo de Whisper en la barra lateral
st.sidebar.header("Configuración del Modelo")
selected_model = st.sidebar.selectbox(
    "Seleccione el modelo de Whisper",
    ("tiny", "base", "small", "medium", "large-v2"),
    index=4  # 'large-v2' por defecto
)

@st.cache_resource
def load_model(model_name):
    """Carga el modelo de Whisper y lo almacena en caché para evitar recargas."""
    logging.info(f"Iniciando la carga del modelo Whisper: {model_name}")
    try:
        model = WhisperModel(
            model_name,
            device=DEVICE,
            compute_type=os.getenv("WHISPER_COMPUTE_TYPE", "float16" if DEVICE == "cuda" else "int8"),
            download_root=os.getenv("WHISPER_CACHE", ".cache/whisper"),
        )
        logging.info("El modelo Whisper se ha cargado correctamente.")
        return model
    except Exception as e:
        st.error(f"Error crítico al cargar el modelo: {e}")
        logging.critical(f"No se pudo cargar el modelo Whisper: {e}", exc_info=True)
        return None

MODEL = load_model(selected_model)

# --- Funciones de Utilidad ---

def _guardar_archivo(uploaded_file, destino: Path) -> None:
    """Guarda un archivo subido por el usuario en el directorio de destino."""
    try:
        with destino.open("wb") as f:
            f.write(uploaded_file.getbuffer())
        logging.info(f"Archivo '{uploaded_file.name}' guardado en '{destino}'")
    except IOError as e:
        logging.error(f"Error de E/S al guardar el archivo {uploaded_file.name}: {e}")
        st.error(f"No se pudo guardar el archivo {uploaded_file.name}.")

def _ffprobe(path: Path) -> Dict[str, Any]:
    """Obtiene metadatos de un archivo multimedia usando ffprobe."""
    cmd = [
        "ffprobe", "-v", "error", "-show_entries", "format=duration",
        "-show_streams", "-print_format", "json", str(path),
    ]
    try:
        res = subprocess.run(cmd, capture_output=True, text=True, check=True)
        return json.loads(res.stdout)
    except (subprocess.CalledProcessError, json.JSONDecodeError) as e:
        logging.error(f"Fallo al ejecutar ffprobe para el archivo {path}: {e}")
        st.error(f"No se pudieron obtener los metadatos del archivo {path.name}.")
        return {}

def _dedupe(text: str, max_rep: int = 3) -> str:
    """Contrae repeticiones de palabras consecutivas para limpiar la transcripción."""
    out, cnt, prev = [], 0, None
    for tok in text.split():
        base = tok.lower()
        if base == prev:
            cnt += 1
            if cnt <= max_rep:
                out.append(tok)
        else:
            prev, cnt = base, 1
            out.append(tok)
    return re.sub(r'\s+', ' ', " ".join(out)).strip()

# --- Funciones de Exportación de Archivos ---

def _export_txt(file_name: str, fecha: str, texto: str) -> bytes:
    """Genera el contenido de un archivo TXT para la descarga."""
    buf = io.BytesIO()
    buf.write(f"Archivo original: {file_name}\nFecha: {fecha}\n\n{texto}".encode("utf-8"))
    buf.seek(0)
    return buf.getvalue()

def _export_docx(file_name: str, fecha: str, texto: str) -> bytes:
    """Genera el contenido de un archivo DOCX para la descarga."""
    doc = Document()
    doc.add_heading("Transcripción de audio", level=0)
    doc.add_paragraph(f"Archivo original: {file_name}")
    doc.add_paragraph(f"Fecha: {fecha}")
    doc.add_paragraph("")
    doc.add_heading("Texto completo", level=1)
    doc.add_paragraph(texto)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# --- Procesamiento de Audio con FFmpeg ---

def _filter_chain_for_classroom() -> str:
    """Cadena de filtros ffmpeg optimizada para grabaciones de voz en aulas."""
    return (
        "highpass=f=120,"          # Elimina ruidos de baja frecuencia (zumbidos, etc.)
        "lowpass=f=7000,"           # Elimina ruidos de muy alta frecuencia
        "afftdn=nr=12:nf=-20,"      # Reducción de ruido de fondo
        "agate=threshold=0.04:ratio=2:attack=20:release=500," # Puerta de ruido para atenuar silencio
        "equalizer=f=3500:width_type=q:width=1:g=3," # Realza la presencia y claridad vocal
        "acompressor=threshold=-18dB:ratio=2.5:attack=20:release=250:makeup=2," # Compresor para nivelar volumen
        "loudnorm=I=-16:TP=-1.5:LRA=11" # Normalización de sonoridad estándar
    )

def _preprocesar_audio(entrada: Path, salida: Path) -> None:
    """Normaliza y convierte un archivo de audio a WAV 16 kHz mono usando FFmpeg."""
    cmd = [
        "ffmpeg", "-y", "-i", str(entrada), "-af", _filter_chain_for_classroom(),
        "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le", str(salida),
    ]
    logging.info(f"Ejecutando comando ffmpeg: {' '.join(cmd)}")
    try:
        subprocess.run(cmd, check=True, capture_output=True, text=True)
        logging.info(f"Preprocesamiento de audio completado para '{entrada.name}'.")
    except subprocess.CalledProcessError as e:
        logging.error(f"ffmpeg falló con el código {e.returncode}: {e.stderr}")
        st.error(f"Error al procesar el audio {entrada.name}. Verifique el log para más detalles.")
        raise

# --- Interfaz Principal de la Aplicación ---

files = st.file_uploader(
    "Seleccione uno o varios archivos de audio (MP3, WAV, M4A)",
    type=["mp3", "wav", "m4a"],
    accept_multiple_files=True,
)

if files:
    try:
        zip_outputs: List[Dict[str, Any]] = []
        fecha_hoy = datetime.now().strftime("%Y-%m-%d")

        for idx, upl in enumerate(files, 1):
            st.header(f"{idx}. {upl.name}")

            if upl.size > MAX_FILE_SIZE:
                st.error("El archivo supera el límite de 1000 MB y será omitido.")
                logging.warning(f"Archivo '{upl.name}' omitido por superar el tamaño máximo.")
                continue

            src_path = Path(UPLOAD_DIR) / upl.name
            _guardar_archivo(upl, src_path)

            wav_path = src_path.with_suffix(".wav")
            with st.spinner("Normalizando y limpiando audio..."):
                _preprocesar_audio(src_path, wav_path)

            meta = _ffprobe(wav_path)
            if not meta: continue # Si ffprobe falla, salta al siguiente archivo.

            dur = float(meta.get("format", {}).get("duration", 0))
            st.info(f"Duración: {dur:.2f}s")

            progress = st.progress(0.0, text="Iniciando transcripción...")
            
            # Variable para controlar la frecuencia de actualización del monitor
            last_update_time = time.time()
            
            try:
                segments_gen, _ = MODEL.transcribe(
                    str(wav_path), language="es", vad_filter=True, beam_size=5,
                )
                segmentos = []
                for seg in segments_gen:
                    segmentos.append(seg)
                    
                    # Actualizar la barra de progreso
                    progress.progress(
                        min(seg.end / dur, 1.0),
                        text=f"Transcribiendo... {min(seg.end / dur, 1.0):.0%}",
                    )

                    # Actualizar el monitor de recursos (máximo 1 vez por segundo)
                    current_time = time.time()
                    if current_time - last_update_time > 1:
                        usage_stats = get_system_usage()
                        gpu_info = f" | GPU: **{usage_stats.get('gpu', 'N/A')}**" if DEVICE == 'cuda' else ""
                        usage_placeholder.info(f"Uso CPU: **{usage_stats['cpu']}** | RAM: **{usage_stats['ram']}**{gpu_info}")
                        last_update_time = current_time

            except RuntimeError as e:
                if "device-side assert triggered" in str(e):
                    st.warning("Fallo de GPU. Reintentando en CPU con un modelo más pequeño...")
                    logging.warning("Fallo de GPU detectado. Cambiando a CPU.")
                    cpu_model = load_model("small")
                    if not cpu_model: raise
                    segments_gen, _ = cpu_model.transcribe(
                        str(wav_path), language="es", vad_filter=True, beam_size=5,
                    )
                    segmentos = list(segments_gen)
                else:
                    raise
            progress.empty()

            if not segmentos:
                st.warning("No se detectó actividad de voz en el audio.")
                logging.warning(f"No se encontró voz en '{upl.name}'.")
                continue

            raw_text = " ".join(s.text for s in segmentos)
            texto_full = _dedupe(raw_text)
            st.success("Transcripción completada.")
            st.text_area("Texto generado", texto_full, height=200, key=f"texto_{idx}")

            with st.expander("Ver segmentos detallados"):
                detalle = "\n".join(f"[{s.start:.2f}s – {s.end:.2f}s] {s.text}" for s in segmentos)
                st.text_area("Segmentos", detalle, height=180, key=f"segmentos_{idx}")

            stem = Path(upl.name).stem
            txt_bytes = _export_txt(upl.name, fecha_hoy, texto_full)
            docx_bytes = _export_docx(upl.name, fecha_hoy, texto_full)

            col1, col2 = st.columns(2)
            with col1:
                st.download_button("Descargar TXT", txt_bytes, file_name=f"{stem}.txt", mime="text/plain", key=f"txt_{idx}")
            with col2:
                st.download_button("Descargar DOCX", docx_bytes, file_name=f"{stem}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"docx_{idx}")

            zip_outputs.append({"name": f"{stem}.txt", "data": txt_bytes})
            zip_outputs.append({"name": f"{stem}.docx", "data": docx_bytes})

        if len(zip_outputs) > 1: # Botón ZIP si hay al menos un archivo procesado
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for item in zip_outputs:
                    zf.writestr(item["name"], item["data"])
            st.download_button(
                "Descargar todo en ZIP", zip_buf.getvalue(),
                file_name="transcripciones.zip", mime="application/zip",
            )

        # Limpieza de archivos de audio temporales
        removed_count = 0
        for f in Path(UPLOAD_DIR).iterdir():
            if f.suffix.lower() in {".mp3", ".wav", ".m4a"}:
                try:
                    f.unlink()
                    removed_count += 1
                except OSError as e:
                    logging.warning(f"No se pudo eliminar el archivo temporal {f}: {e}")
        if removed_count > 0:
            logging.info(f"Se eliminaron {removed_count} archivos de audio temporales.")
            st.toast(f"{removed_count} archivos temporales eliminados.")

    except Exception as exc:
        st.error(f"Error inesperado durante el procesamiento: {exc}")
        logging.critical("Error crítico no manejado en el flujo principal.", exc_info=True)